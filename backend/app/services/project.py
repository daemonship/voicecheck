"""Project service for CRUD and file parsing."""

import os
import tempfile
from typing import Optional
from fastapi import HTTPException, UploadFile
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .supabase import supabase_client
from ..models.project import ProjectCreate, Project


class ProjectService:
    """Handles project operations."""

    def __init__(self):
        self.client = supabase_client

    async def create_project(
        self,
        user_id: str,
        title: str,
        word_count: int,
        chapter_count: int,
        chapters: list[str],
        content_type: str,
        content_ref: str,
        full_text: str = None,
        status: str = "complete",
    ) -> Project:
        """Create a new project in database."""
        try:
            response = self.client.table("projects").insert({
                "user_id": user_id,
                "title": title,
                "word_count": word_count,
                "chapter_count": chapter_count,
                "chapters": chapters,
                "content_type": content_type,
                "content_ref": content_ref,
                "full_text": full_text,
                "status": status,
            }).execute()
            data = response.data[0]
            return Project(**data)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to create project: {str(e)}")

    async def get_project(self, project_id: str, user_id: Optional[str] = None) -> Project:
        """Retrieve a project by ID, optionally checking ownership."""
        try:
            query = self.client.table("projects").select("*").eq("id", project_id)
            response = query.execute()
            if not response.data:
                raise HTTPException(status_code=404, detail="Project not found")
            data = response.data[0]
            if user_id and data["user_id"] != user_id:
                raise HTTPException(status_code=403, detail="Not authorized to access this project")
            return Project(**data)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to retrieve project: {str(e)}")

    async def update_project_status(self, project_id: str, status: str) -> None:
        """Update a project's analysis status."""
        self.client.table("projects").update({"status": status}).eq("id", project_id).execute()

    async def parse_docx(self, file: UploadFile) -> dict:
        """Parse .docx file and extract text, chapters, word count."""
        # Check file size (max 10MB)
        file.file.seek(0, os.SEEK_END)
        size = file.file.tell()
        file.file.seek(0)
        if size > 10 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File size exceeds 10MB limit")

        try:
            # Save to temp file for parsing
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                content = await file.read()
                tmp.write(content)
                tmp_path = tmp.name

            # Parse with python-docx
            doc = Document(tmp_path)

            # Extract text from paragraphs
            full_text = "\n".join([para.text for para in doc.paragraphs])

            # Detect chapters (simple heuristic: lines containing "Chapter" or numbered headings)
            chapters = []
            current_chapter = []
            for para in doc.paragraphs:
                text = para.text.strip()
                # Simple chapter detection
                if text.lower().startswith("chapter") or text.lower().startswith("chapter "):
                    if current_chapter:
                        chapters.append("\n".join(current_chapter))
                        current_chapter = []
                current_chapter.append(text)
            if current_chapter:
                chapters.append("\n".join(current_chapter))

            # If no chapters detected, treat entire document as one chapter
            if len(chapters) == 0:
                chapters = [full_text]

            # Word count (rough)
            word_count = len(full_text.split())

            # Title: first non-empty line or filename
            title = os.path.splitext(file.filename)[0]

            # Clean up temp file
            os.unlink(tmp_path)

            return {
                "title": title,
                "word_count": word_count,
                "chapter_count": len(chapters),
                "chapters": chapters,
                "full_text": full_text,
            }
        except PackageNotFoundError:
            raise HTTPException(status_code=422, detail="Invalid or corrupted .docx file")
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Failed to parse .docx: {str(e)}")

    async def parse_text(self, text: str) -> dict:
        """Parse pasted text and extract chapters, word count."""
        # Detect chapters (simple heuristic)
        lines = text.splitlines()
        chapters = []
        current_chapter = []

        for line in lines:
            stripped = line.strip()
            if stripped.lower().startswith("chapter") or stripped.lower().startswith("chapter "):
                if current_chapter:
                    chapters.append("\n".join(current_chapter))
                    current_chapter = []
            current_chapter.append(line)

        if current_chapter:
            chapters.append("\n".join(current_chapter))

        if len(chapters) == 0:
            chapters = [text]

        word_count = len(text.split())
        title = "Pasted manuscript"

        return {
            "title": title,
            "word_count": word_count,
            "chapter_count": len(chapters),
            "chapters": chapters,
            "full_text": text,
        }
